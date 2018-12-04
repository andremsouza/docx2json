import docx
import json
# from itertools import tee


# fp = "../sample-files/Camila Ferezin do Amarante ( Ginastica ritmica ).docx"
# Convert .docx file to a list of paragraph strings
# Removes empty strings
def docxToStrings(fp):
	txt = [para.text for para in docx.Document(fp).paragraphs]
	return list(filter(None, txt))

# Separate bold paragraphs from non-bold paragraphs
# Consecutive bold or nonbold paragraphs are treated as one element
# 	So, they are concatenated into one string, separated by '\n'
# Returns two lists of strings: "Bold" annd "Nonbold" paragraphs
def docxToStringsSepBold(fp):
	bold, nonbold = [], []
	doc = docx.Document(fp)
	start=0

	# Get first element and append
	prev = None
	for n, paragraph in enumerate(doc.paragraphs):
		if paragraph.text.replace(" ", ""): # Not considering empty paragraphs
			for run in paragraph.runs:
				if run.bold:
					bold.append(paragraph.text)
				else:
					nonbold.append(paragraph.text)
				prev = run
				start = n
				break
			break


	# Get remaining elements
	# In each iteration, compare with previous element for bold/nonbold match
	# If mathing, join the two paragraphs into one string
	# Else, append as usual
	for n, paragraph in enumerate(doc.paragraphs, start+1):
		if paragraph.text.replace(" ", ""): # Not considering empty paragraphs
			for run in paragraph.runs:
				if run.bold and prev.bold: # Bold after bold
					bold[-1] = "\n".join([bold[-1], paragraph.text])
				elif not(run.bold or prev.bold): # Nonbold after Nonbold
					nonbold[-1] = "\n".join([nonbold[-1], paragraph.text])
				elif run.bold: # Bold after nonbold
					bold.append(paragraph.text)
				else: # Nonbold after bold
					nonbold.append(paragraph.text)
				prev = run
				break
	return bold, nonbold

if __name__ == '__main__':
	fp = input("Type the filepath to the .docx: ")
	jdoc = {}
	jdoc["text"] = docxToStrings(fp)
	jdoc["bold"], jdoc["nonbold"] = docxToStringsSepBold(fp)

	print("Writing output to ", fp.replace(".docx", ".json"))
	with open(fp.replace(".docx", ".json"), "w+") as f:
		json.dump(jdoc, f, ensure_ascii=False, indent=4)