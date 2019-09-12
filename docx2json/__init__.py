import docx
import json


def __docxToStrings(inputFile, sepBold=True):
    """
    Convert .docx file to a list of paragraph strings
    Removes empty strings
    If sepBold==True
        Separate bold paragraphs from non-bold paragraphs
        Consecutive bold or nonbold paragraphs are treated as one element
        So, they are concatenated into one string, separated by '\n'
    Returns three lists of strings: "text", "bold" and "nonbold" paragraphs
    """
    # Removing empty paragraphs
    text = [
        para.text.replace("\n", "")
        for para in docx.Document(inputFile).paragraphs
        if para.text.replace("\n", "").replace(" ", "")
    ]
    bold, nonbold = [], []

    if sepBold:
        doc = docx.Document(inputFile)
        start = 0

        # Get first element and append
        prev = None
        for n, paragraph in enumerate(doc.paragraphs):
            para = paragraph.text.replace("\n", "")
            # Not considering empty paragraphs
            if paragraph.text.replace(" ", "").replace("\n", ""):
                for run in paragraph.runs:
                    if run.bold:
                        bold.append(para)
                    else:
                        nonbold.append(para)
                    prev = run
                    start = n
                    break
                break

        # Get remaining elements
        # In each iteration, compare with previous element for bold/nonbold
        # match
        # If matching, join the two paragraphs into one string
        # Else, append as usual
        for paragraph in doc.paragraphs[start + 1:]:
            para = paragraph.text.replace("\n", "")
            # Not considering empty paragraphs
            if paragraph.text.replace(" ", "").replace("\n", ""):
                for run in paragraph.runs:
                    if run.bold and prev.bold:  # Bold after bold
                        bold[-1] = " ".join([bold[-1], para])
                    elif not (run.bold or prev.bold):  # Nonbold after Nonbold
                        nonbold[-1] = " ".join([nonbold[-1], para])
                    elif run.bold:  # Bold after nonbold
                        bold.append(para)
                    else:  # Nonbold after bold
                        nonbold.append(para)
                    prev = run
                    break
    return text, bold, nonbold


def convert(inputFile, sepBold=True, withSave=False, outputFile=None):
    """
    Converts .docx file situated in the inputFile path
    if sepBold==True
        separates paragraphs starting or not with bold characters into
        two lists: "bold", "nonbold"
    if withSave==True
        saves output file to outputFile filepath, defaulted at the same
        location as the input file inputFile
    """
    if outputFile is None:
        outputFile = inputFile.replace(".docx", ".json")
    jdoc = {}

    if sepBold:
        jdoc["text"], jdoc["bold"], jdoc["nonbold"] = __docxToStrings(
            inputFile, sepBold)
    else:
        jdoc["text"] = __docxToStrings(inputFile, sepBold)[0]

    if withSave:
        with open(outputFile, "w+") as f:
            json.dump(jdoc, f, ensure_ascii=True, indent=4)
    return json.dumps(jdoc, ensure_ascii=True, indent=4)


if __name__ == '__main__':
    inputFile = input("Type the filepath to the .docx: ")
    outputFile = input(
        "Type the output filepath (empty if you don't want to save): ")
    print(
        convert(inputFile, True,
                len(outputFile) > 0,
                outputFile if len(outputFile) > 0 else None))
